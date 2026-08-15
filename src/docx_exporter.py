"""
DOCX Exporter: Converts Markdown / Structured PADER report into a formatted Microsoft Word document (.docx).
Optimized for high-throughput table generation.
"""
from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def export_markdown_to_docx(markdown_text: str, output_path: str | Path) -> Path:
    """
    Convert PADER report Markdown to a styled Word Document (.docx).

    Args:
        markdown_text: Markdown formatted report text.
        output_path: Target .docx filepath.

    Returns:
        Path to the written DOCX file.
    """
    doc = Document()
    out_file = Path(output_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    lines = markdown_text.splitlines()
    i = 0
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        # Parse table rows
        rows_data = []
        for line in t_lines:
            if line.strip().startswith("|---"):
                continue
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for r_idx, row in enumerate(rows_data):
            for c_idx in range(num_cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text

                # Styling header row
                if r_idx == 0:
                    tcPr = cell._element.get_or_add_tcPr()
                    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>'))
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(9.0)
                else:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9.0)

        doc.add_paragraph()  # Space after table

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Check for Table
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            i += 1
            continue
        elif table_lines:
            flush_table(table_lines)
            table_lines = []

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(18)
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                r.font.bold = True
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for r in p.runs:
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                r.font.bold = True
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                r.font.bold = True
        elif stripped.startswith("> "):
            # Callout / Note box
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(stripped[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])
        elif re.match(r'^\d+\.\s+', stripped):
            match = re.match(r'^\d+\.\s+', stripped)
            p = doc.add_paragraph(style="List Number")
            p.add_run(stripped[match.end():])
        elif stripped == "---":
            # Horizontal divider
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            p.add_run(stripped)

        i += 1

    if table_lines:
        flush_table(table_lines)

    doc.save(str(out_file))
    return out_file
